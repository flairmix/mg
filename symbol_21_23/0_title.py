from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


text_head = 'АО "МОСГАЗ" Управление "Моспромгаз\"\n\
            125212 Москва, Головинское шоссе, 10г, \n\
            Телефон: +7 (495) 660-20-10, Факс:  +7 (495) 708-00-12'

text_name_object = 'МНОГОФУНКЦИОНАЛЬНАЯ КОМПЛЕКСНАЯ ЖИЛАЯ ЗАСТРОЙКА\n\n\
                    ПО АДРЕСУ: МЕЖДУ УЛ.ЛОБАЧЕВСКОГО И ПЛАТФОРМОЙ «МАТВЕЕВСКОЕ»,\n\
                    КВАРТАЛ 2, КОРПУС 1,2 РАЙОН РАМЕНКИ,\n\
                    ЗАПАДНЫЙ АДМИНИСТРАТИВНЫЙ ОКРУГ ГОРОДА МОСКВЫ'
 

text_type_doc = 'ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ'

sections_87 = {
            1 : ["Раздел 1", "Пояснительная записка"],
            2 : ["Раздел 2", "Схема планировочной организации земельного участка"],
            3 : ["Раздел 3", "Архитектурные решения"],
            4 : ["Раздел 4", "Конструктивные и объемно-планировочные решения"],
            5 : ["Раздел 5", "Сведения об инженерном оборудовании, о сетях инженерно-технического обеспечения, перечень инженерно-технических мероприятий, содержание технологических решений"],
                
            51 : ["Подраздел 1", "Система электроснабжения"],
            52 : ["Подраздел 2", "Система водоснабжения"],
            53 : ["Подраздел 3", "Система водоотведения"],
            54 : ["Подраздел 4", "Отопление, вентиляция и кондиционирование воздуха, тепловые сети"],
            55 : ["Подраздел 5", "Сети связи"],
            56 : ["Подраздел 6", "Система газоснабжения"],
            57 : ["Подраздел 7", "Технологические решения"],
                
            6 : ["Раздел 6", "Проект организации строительства"],
            7 : ["Раздел 7", "Проект организации работ по сносу или демонтажу объектов капитального строительства"],
            8 : ["Раздел 8", "Перечень мероприятий по охране окружающей среды"],
            9 : ["Раздел 9", "Мероприятия по обеспечению пожарной безопасности"],
            10 : ["Раздел 10", "Мероприятия по обеспечению доступа инвалидов"],
            101 : ["Раздел 10.1", "Мероприятия по обеспечению соблюдения требований энергетической эффективности и\
            требований оснащенности зданий,строений и сооружений приборами учета используемых энергетических ресурсов"],
            11 : ["Раздел 11", "Смета на строительство объектов капитального строительства"]
            }

text_part = ["Часть 3", "Отопление, вентиляция и тепломеханические решения АИТ"]

text_book = ["Книга 1", "Отопление, вентиляция и тепломеханические решения АИТ. Корпус 1,2"]

cipher_first_part = "0558.41"
text_cipher = cipher_first_part
text_vol = "Том 11111"


text_date = 'Москва \n 2021'

doc_title = Document()


#0 page
doc_title.add_picture('mospromgas.png', width=Inches(5.0))
last_paragraph = doc_title.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

head = doc_title.add_paragraph(text_head)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

name_object = doc_title.add_paragraph(text_name_object)
name_object.alignment = WD_ALIGN_PARAGRAPH.CENTER

type_doc = doc_title.add_paragraph(text_type_doc)
type_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER

section = doc_title.add_paragraph(sections_87[5][0] + "\n" + sections_87[5][1])
section.alignment = WD_ALIGN_PARAGRAPH.CENTER

if True:
    subsection = doc_title.add_paragraph(sections_87[54][0] + "\n" + sections_87[54][1])
    subsection.alignment = WD_ALIGN_PARAGRAPH.CENTER

if text_part != 0:
    part = doc_title.add_paragraph(text_part[0] + "\n" + text_part[1])
    part.alignment = WD_ALIGN_PARAGRAPH.CENTER

book = doc_title.add_paragraph(text_book[0] + "\n" + text_book[1])
book.alignment = WD_ALIGN_PARAGRAPH.CENTER

cipher = doc_title.add_paragraph(text_cipher + "1111111111 ")     #TODO
cipher.alignment = WD_ALIGN_PARAGRAPH.CENTER

vol = doc_title.add_paragraph(text_vol + "1111111111 ")     #TODO
vol.alignment = WD_ALIGN_PARAGRAPH.CENTER


date = doc_title.add_paragraph(text_date)
date.alignment = WD_ALIGN_PARAGRAPH.CENTER





#end of page
# doc_title.add_page_break()
doc_title.save('title_0.docx')